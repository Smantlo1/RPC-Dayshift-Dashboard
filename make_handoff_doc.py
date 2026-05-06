from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WM_BLUE  = RGBColor(0x00, 0x53, 0xE2)
WM_SPARK = RGBColor(0xFF, 0xC2, 0x20)
DARK     = RGBColor(0x1F, 0x27, 0x37)
GRAY     = RGBColor(0x6B, 0x72, 0x80)

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)


def heading(text, level=1, color=WM_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(18 if level == 1 else 13)
    run.font.color.rgb = color
    return p


def body(text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        rest = p.add_run(text)
        rest.font.size = Pt(11)
        rest.font.color.rgb = DARK
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
    return p


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0053E2')
    pBdr.append(bottom)
    pPr.append(pBdr)


def label_value(label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run(label + ': ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = WM_BLUE
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    return p


# ════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after  = Pt(4)
tr = title.add_run('RPC Dayshift Dashboard')
tr.bold = True
tr.font.size = Pt(26)
tr.font.color.rgb = WM_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
sr = sub.add_run('Developer Handoff Guide')
sr.font.size = Pt(14)
sr.font.color.rgb = GRAY
sr.italic = True

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
note.paragraph_format.space_after = Pt(16)
nr = note.add_run('Plain-English guide for whoever takes over this project')
nr.font.size = Pt(10)
nr.font.color.rgb = GRAY

divider()

# ════════════════════════════════════════════════════════════════════════
# WHAT IS THIS?
# ════════════════════════════════════════════════════════════════════════
heading('What Is This?')
body(
    'This is the RPC Daily Operating Dashboard — a tool built for the Retail '
    'Project Coordinator dayshift team. It turns the standard RPC-to-acting-RPM '
    'daily routine into an interactive checklist and tracking system that runs '
    'entirely in a web browser.'
)
body(
    'The goal is simple: reduce cognitive load, keep the team on schedule, '
    'surface blockers early, and make the daily routine easy to follow without '
    'having to remember everything from scratch each morning.'
)

divider()

# ════════════════════════════════════════════════════════════════════════
# WHERE TO FIND EVERYTHING
# ════════════════════════════════════════════════════════════════════════
heading('Where to Find Everything')

heading('Live Dashboard (shareable link)', level=2, color=DARK)
body('Anyone on Walmart VPN or Eagle WiFi can open this link right now in any browser:')
label_value('URL', 'https://puppy.walmart.com/sharing/s0m0660/rpc-dashboard')
body(
    'This link is permanent — it never changes. When updates are published, '
    'users just refresh their browser to see the latest version.',
    space_after=10
)

heading('Source Code (GitHub)', level=2, color=DARK)
body('The full source code lives here:')
label_value('GitHub Repo', 'https://github.com/Smantlo1/RPC-Dayshift-Dashboard')
body(
    'Everything needed to edit, improve, or re-publish the dashboard is in that repo.',
    space_after=10
)

divider()

# ════════════════════════════════════════════════════════════════════════
# THE MOST IMPORTANT THING TO UNDERSTAND
# ════════════════════════════════════════════════════════════════════════
heading('The Most Important Thing to Understand')
body(
    'The entire dashboard is ONE file: live_dashboard.html. '
    'There is no server, no database to set up, no build process, '
    'and no special software required to run it. '
    'You open the file in a browser and it works. '
    'You edit it in any text editor and it still works. '
    'This was done intentionally to make handoffs like this one as painless as possible.'
)

divider()

# ════════════════════════════════════════════════════════════════════════
# HOW TO GET SET UP
# ════════════════════════════════════════════════════════════════════════
heading('How to Get Set Up (First Time Only)')

body('You will need:')
bullet('Git installed on your computer')
bullet('VS Code (or any text editor you prefer)')
bullet('A GitHub account with access to the repo above')

body('Then do this once:', space_after=4)

p = doc.add_paragraph(style='List Number')
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Open a terminal or command prompt on your computer.')
r.font.size = Pt(11); r.font.color.rgb = DARK

p = doc.add_paragraph(style='List Number')
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Run this command to download the code:')
r.font.size = Pt(11); r.font.color.rgb = DARK
code = doc.add_paragraph('    git clone https://github.com/Smantlo1/RPC-Dayshift-Dashboard.git')
code.paragraph_format.space_after = Pt(4)
for run in code.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

p = doc.add_paragraph(style='List Number')
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Open the folder that was created. You will see one main file: live_dashboard.html.')
r.font.size = Pt(11); r.font.color.rgb = DARK

p = doc.add_paragraph(style='List Number')
p.paragraph_format.space_after = Pt(10)
r = p.add_run('Double-click live_dashboard.html to open it in your browser. That is the dashboard.')
r.font.size = Pt(11); r.font.color.rgb = DARK

divider()

# ════════════════════════════════════════════════════════════════════════
# HOW TO MAKE CHANGES
# ════════════════════════════════════════════════════════════════════════
heading('How to Make Changes')
body(
    'Open live_dashboard.html in VS Code. All of the dashboard logic — '
    'the schedule, the tabs, the task lists, everything — is written in plain '
    'JavaScript inside that one file. You do not need to know a framework. '
    'If you can read and edit text, you can make changes to this file.'
)
body('After making your changes:')
bullet('Save the file')
bullet('Open it in your browser to test that it looks and works the way you expect')
bullet('Once happy, save your changes back to GitHub (see next section)')

divider()

# ════════════════════════════════════════════════════════════════════════
# HOW TO SAVE CHANGES BACK TO GITHUB
# ════════════════════════════════════════════════════════════════════════
heading('How to Save Changes Back to GitHub')
body('Run these three commands in your terminal from inside the project folder:')

for cmd, desc in [
    ('git add live_dashboard.html', 'Tells git which file changed'),
    ('git commit -m "Brief description of what you changed"', 'Saves a snapshot with a note'),
    ('git push', 'Sends it up to GitHub so others can see it'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    code_run = p.add_run(f'    {cmd}')
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(10)
    code_run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
    desc_run = p.add_run(f'   ← {desc}')
    desc_run.font.size = Pt(10)
    desc_run.font.color.rgb = GRAY

doc.add_paragraph().paragraph_format.space_after = Pt(6)

divider()

# ════════════════════════════════════════════════════════════════════════
# HOW TO PUBLISH UPDATES TO THE LIVE LINK
# ════════════════════════════════════════════════════════════════════════
heading('How to Publish Updates to the Live Dashboard Link')
body(
    'Making changes in VS Code and pushing to GitHub updates the code, '
    'but it does not automatically update what people see at the shareable link. '
    'To push your changes to the live link, you upload the file through a tool '
    'called share-puppy, which is built into Code Puppy (the AI tool used to build this).'
)
body('To publish an update:')
bullet(
    'Open Code Puppy and say: ',
    bold_prefix=None
)
# rewrite the bullet to have italic instruction
p = doc.add_paragraph(style='List Bullet')
p.paragraph_format.space_after = Pt(4)
r1 = p.add_run('Open Code Puppy and say: ')
r1.font.size = Pt(11); r1.font.color.rgb = DARK
r2 = p.add_run('"Upload live_dashboard.html to share-puppy with name: rpc-dashboard and business: s0m0660"')
r2.font.size = Pt(11); r2.italic = True; r2.font.color.rgb = WM_BLUE

bullet('The live link updates automatically — the URL stays exactly the same, users just refresh')
bullet('No one needs to be given a new link — it is always the same address')

divider()

# ════════════════════════════════════════════════════════════════════════
# QUICK REFERENCE
# ════════════════════════════════════════════════════════════════════════
heading('Quick Reference')

label_value('Live dashboard link', 'https://puppy.walmart.com/sharing/s0m0660/rpc-dashboard')
label_value('GitHub repo',         'https://github.com/Smantlo1/RPC-Dayshift-Dashboard')
label_value('Main file to edit',   'live_dashboard.html')
label_value('Built with',          'Plain HTML + JavaScript + Tailwind CSS (all in one file)')
label_value('Original owner',      'Skyler (s0m0660@walmart.com)')
label_value('Built by',            'Code Puppy — https://puppy.walmart.com')

divider()

# ════════════════════════════════════════════════════════════════════════
# IF SOMETHING BREAKS
# ════════════════════════════════════════════════════════════════════════
heading('If Something Breaks')
body(
    'Because the whole thing is one file with no external dependencies, '
    'most problems are simple to fix. Here are the most common situations:'
)
bullet(
    'The dashboard looks wrong after an edit — open live_dashboard.html in your browser and check the browser console (press F12) for any error messages. They will tell you exactly what line has a problem.',
)
bullet(
    'Data seems to have disappeared — all data is saved in browser localStorage, which means it is specific to each person\'s browser. It has not been deleted; it is just local to that browser.',
)
bullet(
    'The live link is not updating — make sure you uploaded the latest saved version of live_dashboard.html through share-puppy, not an older copy.',
)
bullet(
    'Need to roll back to a previous version — every change is saved in GitHub history. Go to the repo, click "Commits", find the version you want, and download that file.',
)

divider()

# ════════════════════════════════════════════════════════════════════════
# FOOTER NOTE
# ════════════════════════════════════════════════════════════════════════
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.paragraph_format.space_before = Pt(20)
fr = footer.add_run('Built with Code Puppy · puppy.walmart.com · Questions? Reach out to Skyler (s0m0660)')
fr.font.size = Pt(9)
fr.font.color.rgb = GRAY
fr.italic = True

# ── Save ─────────────────────────────────────────────────────────────────────
out = 'RPC_Dashboard_Handoff_Guide.docx'
doc.save(out)
print(f'Saved: {out}')
